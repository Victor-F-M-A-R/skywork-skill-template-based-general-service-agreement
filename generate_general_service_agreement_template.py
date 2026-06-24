from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TEXT_BLACK = RGBColor(0, 0, 0)
PAGE_WHITE = "FFFFFF"
HEADER_GRAY = "F2F2F2"
BORDER_GRAY = "A6A6A6"


def set_font(run, size=10.5, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = TEXT_BLACK


def shade_element(element, fill=PAGE_WHITE):
    props = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill=PAGE_WHITE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_cell(cell, fill=PAGE_WHITE):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, size="8", color=BORDER_GRAY):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_page_background(document):
    settings = document.settings._element
    background = settings.find(qn("w:displayBackgroundShape"))
    if background is None:
        background = OxmlElement("w:displayBackgroundShape")
        settings.append(background)

    document_elm = document._element
    bg = document_elm.find(qn("w:background"))
    if bg is None:
        bg = OxmlElement("w:background")
        document_elm.insert(0, bg)
    bg.set(qn("w:color"), PAGE_WHITE)


def configure_document(title):
    doc = Document()
    set_page_background(doc)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_BLACK

    add_title(doc, title)
    return doc


def paragraph(doc, text="", size=10.5, bold=False, italic=False, align=None, before=0, after=6, left=0, first_line=0):
    p = doc.add_paragraph()
    shade_paragraph(p)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.left_indent = Cm(left)
    p.paragraph_format.first_line_indent = Cm(first_line)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_title(doc, title):
    paragraph(doc, title, size=17, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    paragraph(doc, "Professional Contract Template", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)


def add_intro(doc, agreement_name):
    paragraph(
        doc,
        f"This {agreement_name} (the \"Agreement\") is entered into as of [EFFECTIVE DATE] (the \"Effective Date\") by and between:",
    )
    paragraph(doc, "[PARTY A LEGAL NAME], a [ENTITY TYPE] organized under the laws of [JURISDICTION], with its principal office at [ADDRESS] (\"Party A\"); and")
    paragraph(doc, "[PARTY B LEGAL NAME], a [ENTITY TYPE] organized under the laws of [JURISDICTION], with its principal office at [ADDRESS] (\"Party B\").")
    paragraph(doc, "Party A and Party B may be referred to individually as a \"Party\" and collectively as the \"Parties\".")


def heading(doc, number, title):
    paragraph(doc, f"{number}. {title}", size=12, bold=True, before=10, after=4)


def clause(doc, number, text):
    paragraph(doc, f"{number} {text}", left=0.78, first_line=-0.78, after=5)


def bullet(doc, text):
    paragraph(doc, f"- {text}", left=0.55, first_line=-0.35, after=4)


def add_clauses(doc, section_number, title, clauses):
    heading(doc, section_number, title)
    for idx, text in enumerate(clauses, 1):
        clause(doc, f"{section_number}.{idx}", text)


def add_page_break_title(doc, title):
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    paragraph(doc, title, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        shade_cell(cell, HEADER_GRAY)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        shade_paragraph(p, HEADER_GRAY)
        r = p.add_run(header)
        set_font(r, bold=True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            shade_cell(cell, PAGE_WHITE)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            shade_paragraph(p)
            r = p.add_run(value)
            set_font(r, size=9.5)
    return tbl


def signature_page(doc, left_label="PARTY A", right_label="PARTY B"):
    add_page_break_title(doc, "SIGNATURE PAGE")
    paragraph(doc, "The Parties have executed this Agreement through their duly authorized representatives as of the Effective Date.")
    rows = [
        (left_label, right_label),
        ("Legal name: [LEGAL NAME]", "Legal name: [LEGAL NAME]"),
        ("By: ______________________________", "By: ______________________________"),
        ("Name: [AUTHORIZED SIGNATORY]", "Name: [AUTHORIZED SIGNATORY]"),
        ("Title: [TITLE]", "Title: [TITLE]"),
        ("Date: [DATE]", "Date: [DATE]"),
    ]
    table(doc, ("", ""), rows)


def schedule(doc, title, rows):
    add_page_break_title(doc, title)
    table(doc, ("Field", "Template Entry"), rows)


def tbd_list(doc, items):
    add_page_break_title(doc, "KEY INFORMATION TO COMPLETE")
    for item in items:
        bullet(doc, item)


def save_doc(doc, output_path):
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))




OUTPUT = Path(__file__).with_name("general_service_agreement_template.docx")


def build():
    doc = configure_document("GENERAL SERVICE AGREEMENT")
    add_intro(doc, "General Service Agreement")

    add_clauses(doc, 1, "Background and Purpose", [
        "The Client desires to engage the Service Provider to provide certain services, deliverables, support, implementation, maintenance, or operational work as described in this Agreement and the applicable Statement of Work.",
        "The Service Provider has represented that it has the experience, personnel, resources, tools, and authority necessary to perform the Services.",
        "This Agreement sets out the general legal and commercial terms that govern the Services and any Statement of Work executed by the Parties.",
    ])
    add_clauses(doc, 2, "Definitions", [
        "\"Services\" means the services, tasks, work, support, implementation, maintenance, or other activities described in a Statement of Work.",
        "\"Deliverables\" means all work product, reports, documents, data, materials, software, content, outputs, or other items to be delivered by the Service Provider.",
        "\"Statement of Work\" or \"SOW\" means a written document signed or approved by both Parties describing project-specific Services, Deliverables, timelines, fees, acceptance criteria, and special terms.",
        "\"Client Materials\" means materials, data, information, systems, instructions, specifications, or assets provided by or on behalf of the Client.",
        "\"Provider Materials\" means the Service Provider's pre-existing tools, templates, methods, know-how, software, documentation, workflows, and background intellectual property.",
    ])
    add_clauses(doc, 3, "Engagement and Scope of Services", [
        "The Client engages the Service Provider, and the Service Provider agrees to provide the Services described in each applicable SOW.",
        "The Services do not include any work, deliverable, support, integration, travel, procurement, licensing, or other obligation not expressly described in this Agreement or an applicable SOW.",
        "If there is a conflict between this Agreement and a SOW, this Agreement controls unless the SOW expressly states that it overrides a specific section of this Agreement.",
    ])
    add_clauses(doc, 4, "Statements of Work", [
        "Each SOW must identify the project name, Services, Deliverables, timeline, milestones, fees, Client dependencies, acceptance criteria, and special terms.",
        "No SOW is binding unless signed by both Parties or otherwise approved in the manner specified in this Agreement.",
        "The Service Provider is not required to begin work under a SOW until all required approvals, deposits, purchase orders, access rights, and Client Materials have been provided.",
    ])
    add_clauses(doc, 5, "Client Responsibilities", [
        "The Client shall provide accurate and complete information, access, personnel, approvals, feedback, systems, and materials reasonably required for the Service Provider to perform the Services.",
        "The Client shall review Deliverables within the applicable review period and provide timely, specific, and written feedback.",
        "The Client is responsible for all business decisions, implementation decisions, and approvals made based on the Services or Deliverables.",
    ])
    add_clauses(doc, 6, "Service Provider Responsibilities", [
        "The Service Provider shall perform the Services in a professional and workmanlike manner using appropriately qualified personnel.",
        "The Service Provider shall comply with applicable laws and regulations in performing the Services.",
        "The Service Provider shall promptly notify the Client of any issue reasonably likely to affect scope, timeline, quality, cost, confidentiality, data security, or legal compliance.",
    ])
    add_clauses(doc, 7, "Deliverables and Acceptance", [
        "The Service Provider shall deliver the Deliverables described in the applicable SOW in the format and by the dates specified in that SOW.",
        "The Client shall inspect or review each Deliverable within [NUMBER] business days after receipt and either accept it or provide a written rejection describing the material non-conformity.",
        "If a Deliverable materially fails to conform to the applicable acceptance criteria, the Service Provider shall use commercially reasonable efforts to correct and resubmit it within [NUMBER] business days.",
        "Unless otherwise stated in the SOW, acceptance does not waive claims for latent defects, confidentiality breaches, intellectual property infringement, or fraud.",
    ])
    add_clauses(doc, 8, "Change Control", [
        "Any change to scope, Deliverables, assumptions, dependencies, timeline, acceptance criteria, fees, or resources must be documented in a written change order approved by both Parties.",
        "The Service Provider is not obligated to perform out-of-scope work unless the Parties agree on the corresponding adjustment to fees, timeline, and responsibilities.",
    ])
    add_clauses(doc, 9, "Fees, Expenses, Invoicing, and Taxes", [
        "The Client shall pay the fees set out in the applicable SOW or fee schedule.",
        "Unless otherwise stated, invoices are payable within [NUMBER] days after receipt of a valid invoice.",
        "Pre-approved out-of-pocket expenses shall be reimbursed at cost if supported by reasonable documentation.",
        "Fees are [inclusive/exclusive] of applicable taxes. Each Party is responsible for taxes imposed on it under applicable law.",
        "The Client shall not withhold undisputed amounts, but may reasonably dispute an invoice in good faith by providing written notice of the disputed portion.",
    ])
    add_clauses(doc, 10, "Term and Renewal", [
        "This Agreement begins on the Effective Date and continues until [END DATE] unless earlier terminated in accordance with this Agreement.",
        "Each SOW has the term stated in that SOW. Expiration or termination of one SOW does not automatically terminate this Agreement or any other SOW.",
    ])
    add_clauses(doc, 11, "Confidentiality", [
        "Each Party shall protect the other Party's Confidential Information and use it only for purposes of performing or receiving the Services.",
        "Confidential Information may be disclosed only to personnel, affiliates, contractors, and advisors with a need to know and confidentiality obligations at least as protective as this Agreement.",
        "These confidentiality obligations survive for [NUMBER] years after disclosure, and trade secrets remain protected for as long as they remain trade secrets under applicable law.",
    ])
    add_clauses(doc, 12, "Intellectual Property", [
        "Each Party retains ownership of its pre-existing intellectual property, tools, templates, know-how, data, and materials.",
        "Ownership or license rights in Deliverables shall be as stated in the applicable SOW. If not stated, the Client receives a non-exclusive, worldwide, perpetual license to use accepted Deliverables for its internal business purposes after full payment.",
        "The Service Provider retains ownership of Provider Materials, methods, templates, reusable components, know-how, and general skills developed or used in performing the Services.",
        "Client Materials remain owned by the Client. The Client grants the Service Provider a limited license to use Client Materials solely to perform the Services.",
    ])
    add_clauses(doc, 13, "Data Protection and Security", [
        "If the Services involve personal data, confidential data, regulated data, or system access, the Parties shall comply with the data protection and security requirements set out in the applicable SOW or data processing schedule.",
        "The Service Provider shall use commercially reasonable administrative, technical, and physical safeguards appropriate to the nature of the information handled.",
    ])
    add_clauses(doc, 14, "Representations and Warranties", [
        "Each Party represents that it has the authority to enter into and perform this Agreement.",
        "The Service Provider warrants that the Services will be performed in a professional and workmanlike manner and that Deliverables will materially conform to the applicable SOW.",
        "Except as expressly stated in this Agreement, all warranties are disclaimed to the maximum extent permitted by law.",
    ])
    add_clauses(doc, 15, "Indemnification", [
        "The Service Provider shall indemnify the Client against third-party claims arising from the Service Provider's gross negligence, willful misconduct, violation of law, or infringement by Deliverables of third-party intellectual property rights.",
        "The Client shall indemnify the Service Provider against third-party claims arising from Client Materials, Client instructions, or the Client's misuse of Deliverables outside the agreed scope.",
    ])
    add_clauses(doc, 16, "Limitation of Liability", [
        "Except for excluded claims, each Party's aggregate liability under this Agreement shall not exceed [LIABILITY CAP].",
        "Neither Party is liable for indirect, incidental, special, punitive, exemplary, or consequential damages, or lost profits, lost revenue, or loss of goodwill, except to the extent not permitted by law.",
        "Excluded claims are: payment obligations, confidentiality breach, fraud, willful misconduct, gross negligence, IP infringement indemnity, data breach, or other claims stated here: [TBD].",
    ])
    add_clauses(doc, 17, "Termination", [
        "Either Party may terminate this Agreement or any SOW for material breach if the breaching Party fails to cure the breach within [NUMBER] days after written notice.",
        "Either Party may terminate this Agreement immediately if the other Party becomes insolvent, ceases business, violates applicable law in a material way, or engages in fraud or willful misconduct.",
        "Termination for convenience is [PERMITTED / NOT PERMITTED]. If permitted, the notice period is [NUMBER] days.",
    ])
    add_clauses(doc, 18, "Effect of Termination", [
        "Upon termination, the Client shall pay undisputed fees for accepted Deliverables and authorized work performed through the effective termination date.",
        "Each Party shall return or destroy the other Party's Confidential Information as required by this Agreement.",
        "Sections intended to survive include payment, confidentiality, intellectual property, indemnity, limitation of liability, dispute resolution, and any provisions that by their nature should survive.",
    ])
    add_clauses(doc, 19, "Independent Contractor", [
        "The Service Provider is an independent contractor and is not an employee, agent, partner, joint venturer, or legal representative of the Client.",
        "The Service Provider has no authority to bind the Client unless expressly authorized in writing.",
    ])
    add_clauses(doc, 20, "Dispute Resolution and Governing Law", [
        "This Agreement is governed by the laws of [GOVERNING LAW].",
        "The Parties shall first attempt to resolve disputes through good-faith negotiation between senior representatives.",
        "If not resolved within [NUMBER] days, disputes shall be resolved by [COURT / ARBITRATION] in [VENUE], in the language [LANGUAGE].",
    ])
    add_clauses(doc, 21, "Miscellaneous", [
        "Neither Party may assign this Agreement without the other Party's prior written consent, except to an affiliate or successor that assumes the assigning Party's obligations.",
        "Any amendment or waiver must be in writing and approved by both Parties.",
        "This Agreement, together with all SOWs and schedules, constitutes the entire agreement between the Parties regarding its subject matter.",
        "This Agreement may be executed in counterparts and by electronic signature where permitted by law.",
    ])

    signature_page(doc, "CLIENT", "SERVICE PROVIDER")
    schedule(doc, "SCHEDULE 1 - STATEMENT OF WORK", [
        ("Project name", "[TBD]"),
        ("Services", "[TBD]"),
        ("Deliverables", "[TBD]"),
        ("Timeline and milestones", "[TBD]"),
        ("Client dependencies", "[TBD]"),
        ("Acceptance criteria", "[TBD]"),
        ("Excluded services", "[TBD]"),
        ("Special terms", "[TBD]"),
    ])
    schedule(doc, "SCHEDULE 2 - FEES AND PAYMENT", [
        ("Fee model", "[Fixed fee / hourly / milestone / retainer / other]"),
        ("Total fees or rates", "[TBD]"),
        ("Invoice timing", "[TBD]"),
        ("Payment deadline", "[TBD]"),
        ("Expenses", "[TBD]"),
        ("Taxes", "[TBD]"),
    ])
    schedule(doc, "SCHEDULE 3 - ACCEPTANCE / SLA", [
        ("Review period", "[TBD]"),
        ("Acceptance criteria", "[TBD]"),
        ("Correction process", "[TBD]"),
        ("Service levels", "[TBD / Not applicable]"),
        ("Service credits", "[TBD / Not applicable]"),
    ])
    tbd_list(doc, [
        "Party legal names, addresses, entity details, and signatories.",
        "Services, Deliverables, SOW details, assumptions, and exclusions.",
        "Fee model, invoice timing, payment deadline, expenses, and taxes.",
        "Acceptance process, review period, correction period, and service levels.",
        "IP ownership/license model, data security terms, liability cap, governing law, and venue.",
    ])
    save_doc(doc, OUTPUT)


if __name__ == "__main__":
    build()
